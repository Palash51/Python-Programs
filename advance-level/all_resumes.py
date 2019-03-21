import os
import sys
import django
import re
import pandas as pd
import xlsxwriter
import textract

from docx import Document
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import StringIO

import nltk, subprocess

sys.path.append("/home/palash/exadatum/ost/osttalent/")
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "osttalent.settings")
django.setup()


# def candidate_email(lst):
# 	"""get email ids of a candidate"""
# 	candidate_email_ids = [i for i in lst if i.__contains__("@")]

# 	all_email_ids = []
# 	if candidate_email_ids:
# 		for candidate in candidate_email_ids:
# 			if candidate.__contains__("@"):
# 				email_ids = re.findall(r"[a-z0-9\.\-+_]+@[a-z0-9\.\-+_]+\.[a-z]+", candidate.lower())
# 				if not email_ids:
# 					email_ids = re.findall('\S+@\S+',candidate.lower())

# 				if email_ids[0] not in all_email_ids:
# 					all_email_ids.append(email_ids[0])

# 	return all_email_ids





def convertDocxToText(path):
	document = Document(path)
	text = "\n".join([para.text for para in document.paragraphs])
	if not text.__contains__("@"):
		return  textract.process(path).decode("utf-8")

	return "\n".join([para.text for para in document.paragraphs])


def convertPDFToText(path):
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    codec = 'utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    fp = open(path, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos=set()
    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password,caching=caching, check_extractable=True):
        interpreter.process_page(page)
    fp.close()
    device.close()
    string = retstr.getvalue()
    retstr.close()
    return string



def getEmail(inputString): 
        '''
        Given an input string, returns possible matches for emails. Uses regular expression based matching.
        '''
        # inputString = inputString.decode(encoding="ascii", errors="ignore")
        email = None
        try:
            pattern = re.compile(r'\S*@\S*')
            matches = pattern.findall(inputString) # Gets all email addresses as a list
            email = matches
        except Exception as e:
            print(e)

        return email



def readFile(fileName):
        '''
        Read a file given its name as a string.
        Modules required: os
        UNIX packages required: antiword, ps2ascii
        '''
        extension = fileName.split(".")[-1]
        if extension == "txt":
            f = open(fileName, 'r')
            string = f.read()
            f.close() 
            return string#, extension
        elif extension == "doc":
            # Run a shell command and store the output as a string
            # Antiword is used for extracting data out of Word docs. Does not work with docx, pdf etc.
            return subprocess.Popen(['antiword', fileName], stdout=subprocess.PIPE, stderr=subprocess.PIPE).communicate()[0]#, extension
        elif extension == "docx":
            try:
                return convertDocxToText(fileName)#, extension
            except:
                return ''
                pass
        
        elif extension == "pdf":
            # ps2ascii converst pdf to ascii text
            # May have a potential formatting loss for unicode characters
            # return os.system(("ps2ascii %s") (fileName))
            try:
                return convertPDFToText(fileName)#, extension
            except:
                return ''
                pass
        else:
            print('Unsupported format')
            return '', ''





def main():
	"""folder path with all resumes"""
	path = '/home/palash/exadatum/ost/all/'
	folder = os.fsencode(path)
	all_resumes = []
	for file in os.listdir(folder):
		filename = os.fsdecode(file)
		if filename.endswith(('.pdf', '.doc', '.docx')):
			all_resumes.append(filename)
	
	# print(all_resumes)
	all_file_email = []
	all_file_content = []
	all_file_email_name = []
	all_file_name = []

	for resume in all_resumes:
		f = open(path + resume, 'r+')
		text = readFile(f.name)
		if not isinstance(text, str):
			text = text.decode(encoding="ascii", errors="ignore") 
		email_ids = getEmail(text)
		all_file_email.append(", ".join(email_ids))
		file_content = text.replace("\n", " ").replace("\x0c", " ").replace("•", "").replace("    ", "").replace("|","").strip()
		all_file_content.append(file_content)
		all_file_name.append(resume)



		file_name_email_id = ""
		for email_id in all_file_email:
			for i in list(email_id):
				if i in list(resume):
					file_name_email_id = file_name_email_id + i

		all_file_email_name.append(file_name_email_id)

	
	data = [
			# {'file_name' : all_file_email_name},
			{'file_name' : all_file_name},
			{'file_email': all_file_email},
			{'file_content' : all_file_content}
		]


	# print(data)
	# print(all_file_email_name)
	final_df = pd.DataFrame()
	for id in range(0,len(data)):
		df = pd.DataFrame.from_dict(data[id])
		final_df = pd.concat([final_df, df], axis=1)

	final_df.to_excel('ostdata.xlsx')



if __name__ == '__main__':
	main()
